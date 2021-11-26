# coding=utf-8
import os
import random
import math
import calendar   #引入日历包
import time
import ast
import docx
import openpyxl
from datetime import datetime
from sendMail import readConfig
from sendMail import deleteSpaceStr
from excelProcess import excelProcess
from sendMail import sendMail
from userInfo import  tableInfo
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from email.mime.multipart import MIMEMultipart

data_dict = {}  #数据字典
month_data = {}  #排班数据字典
curr_year = 0   #当前年份
curr_month = 0  #当前月份
curr_week = 0    #当前星期
table_src = r"./test.xlsx"  #表格路径
table_nam = "Sheet1"   #表格名称
table_ins:excelProcess = excelProcess()  #表格实体

#创建数据字典
def createUserDict():
    index = 1
    max_row = table_ins.max_row
    max_col = table_ins.max_col
    for inx in range(3, max_row):
        user:tableInfo = tableInfo()
        data_dict[str(index)] = user
        rows:list = table_ins.getTableRowValue(inx)
        user.setID(int(rows[0]))
        user.student_name = rows[1]
        user.parent_name = rows[2]
        user.phone = rows[3]
        user.mail = rows[4]
        user.dutys = rows[5:10]  #截取列表，获取值班信息
        user.duty_total = 0
        user.createDutyTab()
        index += 1
    #print(data_dict)

#获取可在当天值班的家长
def getCanDuty(week:int)->list:
    canweeks = []
    for value in data_dict.values():
        data:tableInfo = value
        can = data.duty_tab[str(week)]
        #可在本日值班
        if can != -1:
            canweeks.append(data)
    return  canweeks

#分配值班家长
#num->分配数量  #diss->分配列表
def distributionDuty(num:int, diss:list)->list:
    total = len(diss)
    dis_list = []
    bucket = createSeedBucket(total)
    if num > total:
        print("分配数量超过可分配总数!")
    else:
        for inx in range(1, num + 1):
            seed = math.floor(random.random() * len(bucket))  #向下取整
            dis_list.append(diss[seed])
            bucket.pop(seed)  #从桶中弹出此序号
    return  dis_list

#创建一个种子桶
def createSeedBucket(num:int)->list:
    index = 1
    bucket = []
    while index < num:
        bucket.append(index)
        index += 1
    return bucket

#返回中文大写周
def getChineseWeek(num:int)->str:
    week = ["周一","周二","周三","周四","周五","周六","周日"]
    if num > len(week):
        return ""
    return week[num]

#生成值班数据
def createMonthDuty(year:int, month:int)->dict:
    duty_data = {}
    month_days = calendar.monthrange(year, month)  #获取生成班表那个月的总天数
    for inx in  range(1, month_days[1]):
        prinstr = ""
        format_str = str(year) + str(month) + str(inx)
        prinstr += format_str + " "
        curr_week = datetime.strptime(format_str, "%Y%m%d").weekday()  # 获得班表月第一天周几数据
        #curr_week += 1  #加一取正常月历表示方式
        if curr_week + 1 == 6 or curr_week + 1 == 7:
            continue
        canweeks = getCanDuty(curr_week + 1)  #获取这天可值班的家长
        diss = distributionDuty(2, canweeks)
        new = {}
        new["duty"] = ""
        new["phone"] = ""
        new["mails"] = ""
        new["date"] = datetime.strptime(format_str, "%Y%m%d").isoformat() #日期
        new["week"] = getChineseWeek(curr_week)  #周几
        new["times"] = str(year) + "年" + str(month) + "月" + str(inx) + "日"
        prinstr += new["week"] + "可值班家长: "
        for dix in range(0, len(diss)):
            if (dix < 2):  #选出两名即可
                user:tableInfo = diss[dix]
                 #值班人(学生名＋父母名)
                new["duty"] += "||" + str(user.student_name + user.parent_name)
                new["phone"] += "||" + str(user.phone)
                new["mails"] += "||"+ (str(user.mail))
            else:
                break
        prinstr += new["duty"] + ""
        print(prinstr)
        duty_data[str(inx)] = new
    return  duty_data

#将排班数据写入表格
def writeDutyDataToExcel(src:str, data:dict, excel:excelProcess):
    raw = []   #写入的数据集合
    keys = ["date", "week", "duty" ,"phone", "mails"]  #按照列表给出的key顺序循环
    for value in data.values():
        line = []
        duty_data:dict = value
        for nam in keys:
            line.append(duty_data[nam])
        raw.append(line)
    excel.writeExcel(src, raw)

#将排班信息数据写入到文件夹下
def writeDutyDataToDisk(src:str, data:dict):
    try:
        fs = open(src, "w", encoding="utf-8")
        fs.write(str(data))
    except FileNotFoundError:
        print("无法找到该路径！")
    fs.close()

#生成docx文档
def createDocument(src:str, duty_data:dict):
    file_name = "邀请函.docx"
    doc_dict = readDocxConfig(src)
    for value in duty_data.values():
        contex = ""
        time_str = value["times"]
        title = "值班" + doc_dict["title"]
        parents:list = (value["duty"].split("||"))[1:3]
        for val in parents:
            new_src = './document/' + (time_str + val + title) + ".docx"  # 拼接文档名称
            contex = "尊敬的 " + val + "\n" + doc_dict["context"]  #邀请函开头
            new_dict = {
                "title": doc_dict["title"],
                "context": contex
            }
            print("生成" + new_src)
            writeDocument(new_src, new_dict)

#写入文档
def writeDocument(src:str, data:dict):
    # 创建文档对象
    document = docx.Document()
    # 设置文档标题，中文要用unicode字符串
    title = document.add_heading(0)
    #title.alignment = WD_PARAGRAPH_ALIGNMENT  #居中
    title_run = title.add_run(u'' + data["title"])  #添加标题内容
    title_run.font.name = u"微软雅黑"
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  #设置标题字体
    title_run.font.size = Pt(24)  #设置标题字体大小
    title_run.font.color.rgb = RGBColor(0, 0, 0)  #字体颜色
    # 往文档中添加段落
    p = document.add_paragraph(data["context"])
    # p.add_run('bold ').bold = True
    # p.add_run('and some ')
    # p.add_run('italic.').italic = True
    # 添加一级标题
    # document.add_heading(u'一级标题, level = 1', level=1)
    # document.add_paragraph('Intense quote', style='IntenseQuote')
    # # 添加无序列表
    # document.add_paragraph('first item in unordered list', style='ListBullet')
    # # 添加有序列表
    # document.add_paragraph('first item in ordered list', style='ListNumber')
    # document.add_paragraph('second item in ordered list', style='ListNumber')
    # document.add_paragraph('third item in ordered list', style='ListNumber')
    # # 添加图片，并指定宽度
    # #document.add_picture('cat.png', width=Inches(2.25))
    # # 添加表格: 1行3列
    # table = document.add_table(rows=1, cols=3)
    # # 获取第一行的单元格列表对象
    # hdr_cells = table.rows[0].cells
    # # 为每一个单元格赋值
    # # 注：值都要为字符串类型
    # hdr_cells[0].text = 'Name'
    # hdr_cells[1].text = 'Age'
    # hdr_cells[2].text = 'Tel'
    # # 为表格添加一行
    # new_cells = table.add_row().cells
    # new_cells[0].text = 'Tom'
    # new_cells[1].text = '19'
    # new_cells[2].text = '12345678'
    # # 添加分页符
    # document.add_page_break()
    # # 往新的一页中添加段落
    # p = document.add_paragraph('This is a paragraph in new page.')
    # 保存文档
    document.save(src)


#值班排班流程
def dutyFunc():
    global curr_year
    global curr_month
    global month_data
    file_name = "月值班表.xlsx"
    duty_file = "./dutyinfo/"
    month = input("请输入要生成的值班月份:\n")
    month_value = int(month)
    locale_time = time.localtime()
    curr_year = locale_time.tm_year  #获取年份
    curr_month = month_value
    table_ins.readExcel(table_src, table_nam)
    createUserDict()
    month_data = createMonthDuty(curr_year, month_value)
    createDocument("", month_data)
    duty_file += str(curr_year) + str(month_value) + "Duty.txt"
    writeDutyDataToDisk(duty_file, month_data)
    writeDutyDataToExcel(str(month_value) + file_name, month_data, table_ins)

#发送邀请函
def sendInvitation():
    global month_data
    curr_day = 0
    curr_year = 0
    locale_time = time.localtime()
    month = input("请输入发送邀请函的所在月日(格式:12|1):[回车默认发送本月本天的邀请函]\n")
    if month == "\n" or month == "":
        curr_day = locale_time.tm_mday
        month_value = locale_time.tm_mon
    else:
        str_arr = month.split("|")
        curr_day = int(str_arr[1])
        month_value = int(str_arr[0])
    curr_year = locale_time.tm_year
    format_str = str(curr_year) + str(month_value) + str(curr_day)
    week = datetime.strptime(format_str, "%Y%m%d").weekday()  # 获得班表月第一天周几数据
    file_name = "./dutyinfo/" + format_str + "Duty.txt"
    try:
        fs = open(file_name, "r", encoding="utf-8")
        dict_str = fs.read()
        month_data = ast.literal_eval(dict_str) #读取本地磁盘的排班数据信息
        fs.close()
        duty_info:dict = getDutyParentInfo(month_data, curr_day)
        if week == 6 or week == 7:
            print("休息日是不用值班的哦！亲")
        else:
            sendToMail(duty_info)
    except FileNotFoundError:
        print("找不到当月值班数据, 请重新运行程序生成！")

#查找出当月当日要值班的家长
#month_data -> 排班数据  day -> 时间
def getDutyParentInfo(month_data:dict, day:int)->list:
    dutys = []
    for key, value in month_data.items():
        data:tableInfo = value
        if int(key) == day:
            return data
    return None

#读取文档配置
def readDocxConfig(src:str)->dict:
    # 文件读取字典
    file_dict = {
        "title": "",
        "context": "",
    }
    ctx_tag = False
    fs = open("./docxconfig.txt", "r", encoding="utf-8")  # 只读打开
    while 1:
        line = fs.readline()
        if not line:
            break
        else:
            #若内容读取标记打开，则执行内容文本链接
            if ctx_tag == True:
                if line.find("=") != -1:    #当读取到的某行带=号，则表示内容读取结束
                    ctx_tag = False
                file_dict["context"] += line
            if line.find("邀请函标题") != -1:
                file_dict["title"] = deleteSpaceStr(line.split("=")[1])
            if line.find("邀请函内容") != -1:
                ctx_tag = True
                file_dict["context"] += line.split("=")[1]
    fs.close()
    return file_dict

#根据排班信息与本地配置发送邮件
def sendToMail(duty:dict):
    mail_str:str = duty["mails"]
    time_str:str = ((duty["date"]).split("T"))[0]  #时间
    date_arr:list = time_str.split("-")  #年月日
    parents:list = ((duty["duty"]).split("||"))[1:3]  #获取值班人员
    mails_arr:list = (mail_str.split("||"))[1:3]  #截取邮箱信息
    mail_config = readConfig("")  #读取config.txt文件
    for val in parents:
        #str(int(date_arr[1]))消除前面带0的字符
        file_name:str = date_arr[0] + "年" + str(int(date_arr[1])) + "月" + str(int(date_arr[2])) + "日"
        file_name += (val + "值班邀请函.docx")
        # file_name = "2021年12月1日苏岩妈妈值班邀请函.docx"
        mail = sendMail()
        mail.m_sender = str(mail_config["m_sender"])
        mail.m_receivers = mails_arr
        #mail.m_receivers = ["cj18760099334@163.com"]
        mail.m_context = str(mail_config["m_context"])
        mail.m_title = str(mail_config["m_title"])
        mail.setMailSmtp(str(mail_config["m_smtp"]))
        mail.setEmpowerCode(str(mail_config["m_code"]))
        mime:MIMEMultipart = mail.setEnclosureMail("./document/", file_name)
        mail.sendMailReady(mime)
        mail.sendMail()

def runMain():
    # doc_dict = readDocxConfig("")
    # writeDocument('./document/demo1.docx', doc_dict)
    select = input("功能选择 (1:生成值班表 (2:发送当天邀请函  [输入数字1 or 2]:\n")
    select_value = int(select)
    if select_value == 1:
        dutyFunc()
    if select_value == 2:
        sendInvitation()

runMain()
os.system("pause")

